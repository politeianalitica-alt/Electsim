"""
Export Service — Bloque 9.

Exporta borradores de informe a distintos formatos:
  export_report_markdown, export_report_html,
  export_report_docx, export_report_pdf.

Degradación graceful:
  - Markdown siempre disponible
  - HTML con Markdown básico si no hay markdown2
  - DOCX solo si python-docx está instalado
  - PDF solo si weasyprint está instalado
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Flags de disponibilidad
try:
    import markdown2 as _md2  # type: ignore
    _MARKDOWN2_OK = True
except ImportError:
    _MARKDOWN2_OK = False

try:
    from docx import Document as _DocxDocument  # type: ignore
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import weasyprint as _weasyprint  # type: ignore
    # WeasyPrint puede importar pero fallar por libs del sistema
    _weasyprint.HTML(string="<p>test</p>").write_pdf()
    _WEASYPRINT_OK = True
except Exception:
    _WEASYPRINT_OK = False


def export_report_markdown(
    report_id: str,
    output_path: str | None = None,
    engine: Any | None = None,
) -> str:
    """
    Exporta un informe como cadena Markdown.

    Args:
        report_id: ID del informe.
        output_path: Ruta opcional donde guardar el archivo .md.
        engine: SQLAlchemy engine.

    Returns:
        Cadena Markdown del informe.
    """
    from services.documents.draft_service import get_report

    report = get_report(report_id, engine=engine)
    if not report:
        logger.warning("export_report_markdown: report %s not found", report_id)
        return f"# Informe no encontrado\n\nID: {report_id}\n"

    lines: list[str] = []

    # Cabecera
    lines.append(f"# {report.title}")
    lines.append("")
    lines.append(f"**Tipo:** {report.report_type}  ")
    lines.append(f"**Estado:** {report.status}  ")
    if report.client_id:
        lines.append(f"**Cliente:** {report.client_id}  ")
    if report.created_at:
        lines.append(f"**Creado:** {report.created_at.strftime('%Y-%m-%d %H:%M UTC')}  ")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Secciones
    for section in report.sections:
        title = section.get("title", "Sección")
        body = section.get("body_markdown", "")
        lines.append(f"## {title}")
        lines.append("")
        lines.append(body)
        lines.append("")

    # Fuentes
    if report.source_objects:
        lines.append("---")
        lines.append("")
        lines.append("## Fuentes")
        lines.append("")
        for obj in report.source_objects:
            lines.append(f"- **{obj.get('type', 'objeto')}**: `{obj.get('id', '')}`")
        lines.append("")

    # Evidencias
    if report.evidence_ids:
        lines.append("## Evidencias")
        lines.append("")
        for eid in report.evidence_ids:
            lines.append(f"- `{eid}`")
        lines.append("")

    md_text = "\n".join(lines)

    if output_path:
        try:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(md_text, encoding="utf-8")
            logger.info("export_report_markdown: saved to %s", output_path)
        except Exception as exc:
            logger.warning("export_report_markdown write: %s", exc)

    return md_text


def export_report_html(
    report_id: str,
    output_path: str | None = None,
    engine: Any | None = None,
) -> str:
    """
    Exporta un informe como HTML.

    Usa markdown2 si está disponible; en caso contrario convierte
    el Markdown con transformaciones básicas de regex.

    Returns:
        Cadena HTML del informe.
    """
    md_text = export_report_markdown(report_id, engine=engine)

    if _MARKDOWN2_OK:
        import markdown2 as md2
        body_html = md2.markdown(
            md_text,
            extras=["tables", "fenced-code-blocks", "header-ids"],
        )
    else:
        body_html = _simple_md_to_html(md_text)

    # Buscar el título para la etiqueta <title>
    title_line = next(
        (l.lstrip("# ").strip() for l in md_text.splitlines() if l.startswith("# ")),
        "Informe"
    )

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title_line}</title>
  <style>
    body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 860px;
            margin: 40px auto; padding: 0 20px; color: #222; line-height: 1.6; }}
    h1 {{ color: #1a1a2e; border-bottom: 2px solid #4a90d9; padding-bottom: 8px; }}
    h2 {{ color: #2c3e50; margin-top: 2em; }}
    h3 {{ color: #34495e; }}
    hr {{ border: none; border-top: 1px solid #ddd; margin: 2em 0; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
    th {{ background: #f5f7fa; }}
    code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
    pre {{ background: #f4f4f4; padding: 12px; border-radius: 4px; overflow-x: auto; }}
    @media print {{
      body {{ margin: 0; padding: 20px; }}
    }}
  </style>
</head>
<body>
{body_html}
</body>
</html>"""

    if output_path:
        try:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(html, encoding="utf-8")
            logger.info("export_report_html: saved to %s", output_path)
        except Exception as exc:
            logger.warning("export_report_html write: %s", exc)

    return html


def export_report_docx(
    report_id: str,
    output_path: str | None = None,
    engine: Any | None = None,
) -> bytes | None:
    """
    Exporta un informe como archivo DOCX (python-docx).

    Returns:
        Bytes del archivo DOCX, o None si python-docx no está disponible.
    """
    if not _DOCX_OK:
        logger.warning("export_report_docx: python-docx no instalado")
        return None

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from services.documents.draft_service import get_report

    report = get_report(report_id, engine=engine)
    if not report:
        logger.warning("export_report_docx: report %s not found", report_id)
        return None

    doc = Document()

    # Estilo del documento
    doc.core_properties.title = report.title
    if report.created_at:
        doc.core_properties.created = report.created_at

    # Título principal
    heading = doc.add_heading(report.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadatos
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Tipo: {report.report_type}  |  Estado: {report.status}")
    if report.created_at:
        meta_para.add_run(f"  |  {report.created_at.strftime('%Y-%m-%d %H:%M UTC')}")
    meta_para.runs[0].font.size = Pt(10)
    meta_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Secciones
    for section in report.sections:
        title = section.get("title", "Sección")
        body = section.get("body_markdown", "")

        doc.add_heading(title, level=1)

        # Convertir Markdown básico a párrafos
        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("1. ") or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == "."):
                p = doc.add_paragraph(stripped[3:] if stripped.startswith("1. ") else stripped[2:],
                                      style="List Number")
            else:
                doc.add_paragraph(stripped)

    # Fuentes
    if report.source_objects:
        doc.add_page_break()
        doc.add_heading("Fuentes", level=1)
        for obj in report.source_objects:
            doc.add_paragraph(
                f"{obj.get('type', 'objeto')}: {obj.get('id', '')}",
                style="List Bullet",
            )

    # Serializar a bytes
    import io
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    if output_path:
        try:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(content)
            logger.info("export_report_docx: saved to %s", output_path)
        except Exception as exc:
            logger.warning("export_report_docx write: %s", exc)

    return content


def export_report_pdf(
    report_id: str,
    output_path: str | None = None,
    engine: Any | None = None,
) -> bytes | None:
    """
    Exporta un informe como PDF usando WeasyPrint.

    Returns:
        Bytes del PDF, o None si WeasyPrint no está disponible.
    """
    if not _WEASYPRINT_OK:
        logger.warning("export_report_pdf: weasyprint no instalado. "
                       "Usa: pip install weasyprint")
        return None

    html_text = export_report_html(report_id, engine=engine)

    try:
        import weasyprint
        pdf_bytes = weasyprint.HTML(string=html_text).write_pdf()

        if output_path:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(pdf_bytes)
            logger.info("export_report_pdf: saved to %s", output_path)

        return pdf_bytes
    except Exception as exc:
        logger.error("export_report_pdf: %s", exc)
        return None


def get_export_capabilities() -> dict:
    """Devuelve qué formatos de exportación están disponibles."""
    return {
        "markdown": True,
        "html": True,
        "html_enhanced": _MARKDOWN2_OK,
        "docx": _DOCX_OK,
        "pdf": _WEASYPRINT_OK,
    }


# ── Helpers ────────────────────────────────────────────────────────────────────

def _simple_md_to_html(md_text: str) -> str:
    """Convierte Markdown básico a HTML sin dependencias externas."""
    import re

    lines = md_text.splitlines()
    html_lines: list[str] = []
    in_list = False

    for line in lines:
        # Headings
        if line.startswith("# "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h1>{_escape_html(line[2:])}</h1>")
        elif line.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{_escape_html(line[3:])}</h2>")
        elif line.startswith("### "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h3>{_escape_html(line[4:])}</h3>")
        elif line.startswith("---"):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<hr>")
        elif line.startswith("- ") or line.startswith("* "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{_inline_md(_escape_html(line[2:]))}</li>")
        elif line.strip() == "":
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<br>")
        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<p>{_inline_md(_escape_html(line))}</p>")

    if in_list:
        html_lines.append("</ul>")

    return "\n".join(html_lines)


def _escape_html(text: str) -> str:
    """Escapa caracteres HTML básicos."""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _inline_md(text: str) -> str:
    """Convierte negritas e itálicas de Markdown en HTML."""
    import re
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text
